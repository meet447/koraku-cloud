"""Build .docx files from a JSON spec (CLI: python -m koraku.artifacts.docx_build)."""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any


def _load_spec(path: str | None, inline: str | None) -> dict[str, Any]:
    if path:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    if inline:
        return json.loads(inline)
    return {}


def build_docx(spec: dict[str, Any], out_path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise SystemExit(
            "python-docx is required. Install with: pip install 'koraku[artifacts]'"
        ) from e

    doc = Document()
    title = str(spec.get("title") or "").strip()
    if title:
        heading = doc.add_heading(title, level=0)
        for run in heading.runs:
            run.font.size = Pt(22)

    subtitle = str(spec.get("subtitle") or "").strip()
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True

    for section in spec.get("sections") or []:
        if not isinstance(section, dict):
            continue
        heading = str(section.get("heading") or section.get("title") or "").strip()
        if heading:
            level = int(section.get("level") or 1)
            doc.add_heading(heading, level=min(max(level, 1), 3))
        body = section.get("body") or section.get("text") or ""
        if isinstance(body, list):
            for item in body:
                text = str(item).strip()
                if text:
                    doc.add_paragraph(text, style="List Bullet")
        elif str(body).strip():
            for para in str(body).split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        for row in section.get("bullets") or []:
            text = str(row).strip()
            if text:
                doc.add_paragraph(text, style="List Bullet")

        table_rows = section.get("table")
        if isinstance(table_rows, list) and table_rows:
            cols = max(len(r) if isinstance(r, list) else 1 for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(table_rows):
                cells = row if isinstance(row, list) else [row]
                for ci in range(cols):
                    val = str(cells[ci]) if ci < len(cells) else ""
                    table.rows[ri].cells[ci].text = val

    out = Path(out_path).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Build a .docx from JSON spec")
    parser.add_argument("--out", required=True, help="Output .docx path")
    parser.add_argument("--spec", help="Path to JSON spec file")
    parser.add_argument("--json", dest="inline_json", help="Inline JSON spec")
    parser.add_argument("--title", help="Quick title when no spec file")
    args = parser.parse_args(argv)

    spec = _load_spec(args.spec, args.inline_json)
    if args.title and not spec.get("title"):
        spec["title"] = args.title

    path = build_docx(spec, args.out)
    print(json.dumps({"ok": True, "path": path, "type": "docx"}))
    return 0


if __name__ == "__main__":
    sys.exit(main())
